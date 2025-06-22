

import io
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_from_pdf(pdf_file):
    """
    Extracts text from an uploaded PDF file.

    Returns:
        A tuple containing:
        - str: The extracted text.
        - str: An error message if something goes wrong, otherwise None.
    """
    try:
        pdf_reader = PdfReader(pdf_file)
        text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
        return text, None
    except Exception as e:
        return None, f"Error reading PDF file: {e}"


def create_docx_from_json(resume_data: dict) -> io.BytesIO:
    """Creates a DOCX file in memory from the generated resume JSON data."""
    document = Document()
    
    # Title (Name and Contact)
    if 'contact' in resume_data and 'name' in resume_data['contact']:
        document.add_heading(resume_data['contact']['name'], level=0)
    
    contact_info = resume_data.get('contact', {})
    contact_line = " | ".join(filter(None, [
        contact_info.get('email'), 
        contact_info.get('phone'), 
        contact_info.get('linkedin')
    ]))
    if contact_line:
        p = document.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    section_map = {
        "summary": "Professional Summary",
        "experience": "Work Experience",
        "education": "Education",
        "skills": "Skills",
        "projects": "Projects"
    }

    for key, title in section_map.items():
        if key in resume_data and resume_data[key]:
            document.add_heading(title, level=1)
            content = resume_data[key]
            
            if isinstance(content, str):
                document.add_paragraph(content)
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, dict): 
                        primary_line = item.get('title', '')
                        if 'company' in item:
                            primary_line += f" at {item.get('company', '')}"
                        elif 'institution' in item:
                             primary_line += f" at {item.get('institution', '')}"
                        
                        p = document.add_paragraph()
                        p.add_run(primary_line).bold = True

                        loc_date_line = " | ".join(filter(None, [item.get('location'), item.get('dates')]))
                        if loc_date_line:
                            p = document.add_paragraph()
                            p.add_run(loc_date_line).italic = True

                        if 'description' in item:
                            description = item['description']
                            if isinstance(description, list):
                                for point in description:
                                    document.add_paragraph(point, style='List Bullet')
                            else:
                                 document.add_paragraph(description, style='List Bullet')
                    
                    elif isinstance(item, str):
                        document.add_paragraph(item, style='List Bullet')

  
    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio